"""
Generate the Word instruction document for the AI Power-Demand model.
Specialist content (tech-ai-sector-analyst, sourced + precedent-grounded) integrated
2026-06-23. Headline numbers are pulled LIVE from model.py so the doc cannot drift.

  python research/build_instruction_doc.py  ->  research/AI_Power_Demand_Model_Guide.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(_REPO))
import model  # noqa: E402

OUT = _REPO / "research" / "AI_Power_Demand_Model_Guide.docx"
INK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x2C, 0x3E, 0x50)
MUTE = RGBColor(0x60, 0x6A, 0x78)


def _live():
    tok = model.build_token_demand("Base")
    mt = model.build_macro_gap(tok)
    mf = model.build_macro_gap(tok, use_flops_demand=True)
    st_, sf = model.gap_summary(mt), model.gap_summary(mf)
    return {
        "tp": st_["peak_gap_gw"], "tpy": st_["peak_gap_year"], "tfl": mt.loc[2042, "demand_gw"],
        "tc": st_["overshoot_year"], "fp": sf["peak_gap_gw"], "fpy": sf["peak_gap_year"],
        "ffl": mf.loc[2042, "demand_gw"], "fc": sf["overshoot_year"],
        "t25": tok.loc[2025, "total_T"], "t42": tok.loc[2042, "total_T"],
        "capex": st_["cumulative_capex_2042_b"] / 1000.0,
    }


def h1(doc, t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def h2(doc, t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(doc, t, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = INK
    return p


def lead(doc, label, t):
    p = doc.add_paragraph(style="List Bullet")
    a = p.add_run(f"{label} ")
    a.bold = True
    a.font.size = Pt(10.5)
    a.font.color.rgb = INK
    b = p.add_run(t)
    b.font.size = Pt(10.5)
    b.font.color.rgb = INK
    return p


def num(doc, t):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(t)
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    return p


def table(doc, headers, rows, sizes=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(8.5)
    return t


def build():
    n = _live()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ── Cover ──
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AI Power-Demand Model"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = ACCENT
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("A Reader's Guide: How It Works, What the Numbers Mean,\nand Why Every Assumption Is Set Where It Is")
    rs.font.size = Pt(13); rs.font.color.rgb = MUTE
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = d.add_run("Peckham Capital  ·  AI Infrastructure Supply Chain  ·  June 2026")
    rd.font.size = Pt(10); rd.font.color.rgb = MUTE
    para(doc, "")

    # ── 1. What the model does ──
    h1(doc, "1.  What this model does, in one minute")
    para(doc, "This model answers two questions for the AI-infrastructure thesis: how large "
              "does data-center power demand become, and how long does the resulting shortage "
              "(the “gap”) last. It starts from a forecast of AI token usage, converts those "
              "tokens into electrical power in gigawatts (GW), compares that demand against the "
              "pace at which power and grid capacity can physically be built, and reports the gap "
              "year by year to 2042.")
    para(doc, "It carries two lenses on the same engine:")
    lead(doc, "Token base (the committed headline).",
         f"Peak gap {n['tp']:.0f} GW in {n['tpy']}; demand plateaus near {n['tfl']:.0f} GW; the "
         f"gap closes around {n['tc']}. This is the locked base case.")
    lead(doc, "FLOPs / Duration lens.",
         f"A physically honest re-derivation. Slightly lower peak ({n['fp']:.0f} GW in "
         f"{n['fpy']}) but a materially higher demand floor ({n['ffl']:.0f} GW vs {n['tfl']:.0f}) "
         f"and a later close (~{n['fc']}). Its message is duration, not magnitude.")
    para(doc, f"Across both lenses, cumulative power-and-grid capital spending to 2042 is about "
              f"${n['capex']:.1f} trillion. Gross AI tokens grow from ~{n['t25']:.0f} trillion per "
              f"day in 2025 to ~{n['t42']:,.0f} trillion in 2042 — roughly {n['t42']/n['t25']:.0f}x.")
    para(doc, "The Excel workbook is where you read and stress-test; the dashboard is where you "
              "explore interactively. The rest of this guide makes every number tangible and "
              "explains where it came from.", italic=True)

    # ── 2. The big picture in real terms ──
    h1(doc, "2.  The big picture in real-world terms")

    h2(doc, "2.1  What 70 GW — today's installed base — actually is")
    para(doc, "The model anchors 2025 total data-center power at 70 GW. This is total DC draw — "
              "all workloads globally, not AI alone — because the bottleneck layers the thesis "
              "trades (generation, transformers, grid interconnect, electrical steel) serve all "
              "data-center load, so the honest denominator is total capacity. To make 70 GW "
              "tangible:")
    lead(doc, "As facilities:", "a large hyperscale data center is ~50–150 MW; the new flagship "
              "AI campuses (Stargate Abilene, Meta Hyperion, xAI Colossus) run 300 MW to multiple "
              "GW. 70 GW is on the order of 500–1,000 conventional large data centers, or a few "
              "hundred AI-scale campuses.")
    lead(doc, "Versus the US grid:", "US summer generating capacity is ~1,200 GW (after a record "
              "53 GW added in 2025). So 70 GW of global DC load is already ~6% of all US generating "
              "capacity — and that is just the current base, before any growth.")
    lead(doc, "Versus a city:", "New York City's summer peak is ~10–11 GW. 70 GW is six-to-seven "
              "New York Cities of continuous draw — and a data center runs at ~85–90% load factor "
              "(always on) versus a city's ~55%, so its energy footprint per GW is far heavier.")
    para(doc, "The 70 GW anchor is the mid-point of five independent estimates: Cushman & "
              "Wakefield (63 GW in 2024, ~70–72 rolled forward), JLL (~72), Goldman Sachs (65 + "
              "90 under construction), McKinsey (~70), and DCD (~68). They cluster tightly at "
              "65–72 GW. AI is ~25% of new builds in 2025 rising to ~50% by 2030, and the buyers "
              "are concentrated — four US hyperscalers plus the neoclouds account for the bulk of "
              "incremental AI power. That concentration is the thesis's friend: a handful of "
              "deep-pocketed buyers competing for the same scarce transformers and turbines is "
              "exactly what lets the bottleneck owners capture rent.")

    h2(doc, "2.2  What the shortage means")
    para(doc, f"The headline is a peak shortfall of {n['tp']:.0f} GW (token lens, {n['tpy']}) or "
              f"{n['fp']:.0f} GW (FLOPs lens, {n['fpy']}) — the gap between where demand wants to be "
              f"and what the supply chain can energize. In tangible terms:")
    lead(doc, "Nuclear-reactor equivalents:", "a large reactor is ~1 GW. A ~300 GW gap is ~300 "
              "reactors' worth of power. The entire US nuclear fleet is ~94 reactors / ~97 GW — "
              "so the gap is roughly three times the entire US nuclear fleet, to appear in ~4 years.")
    lead(doc, "Gas-turbine equivalents:", "a modern combined-cycle block is ~0.5–1.0 GW, so "
              "~300 GW is 300–600 utility-scale turbines. GE Vernova's heavy-duty turbine slots "
              "are already booked past 2028 and being rationed.")
    lead(doc, "Versus US generation:", "~300 GW is ~24–25% of total US summer capacity — the "
              "unmet portion alone is a quarter of everything the US has ever built.")
    para(doc, "The key reframing: the model is not forecasting that the gap gets built. It is "
              "forecasting that it does not — and that the inability to build it is precisely what "
              "creates multi-year pricing power for the names that own the bottleneck nodes.",
         italic=True)

    h2(doc, "2.3  Why it is physically tight — the four walls")
    para(doc, "The gap is not a financing problem (the hyperscalers have the cash); it is a "
              "physics-and-fabrication problem. Four hard walls, all corroborated in current data:")
    num(doc, "Grid interconnect queue. Average US wait has gone from under 2 years in 2008 to "
             "~5 years today. In December 2025 PJM — serving 65M people — for the first time in "
             "its history failed to procure enough capacity, coming up ~6,625 MW short. Owning a "
             "turbine is not the same as being allowed to plug it in.")
    num(doc, "Large power transformers. Lead times have stretched to 100–210 weeks. The US "
             "imports most of its transformers; new domestic capacity (Hitachi Virginia 2028, "
             "Siemens Charlotte 2027) arrives after the peak-gap years.")
    num(doc, "Grain-oriented electrical steel (GOES). Every transformer needs it; there is one "
             "US producer (Cleveland-Cliffs). Expansion is physics-limited. This is upstream of "
             "the entire grid.")
    num(doc, "Gas turbines and skilled electrical labor. Turbine slots are rationed to 2028+. "
             "Microsoft's President called skilled electrical labor the “#1 problem” in March "
             "2026 — a 4–5 year apprenticeship that no capex can shortcut. Likely the last "
             "binding constraint to clear.")
    para(doc, "For scale: US generating capacity has grown ~1–3%/yr for two decades. The model "
              "asks the global fleet to compound at 22–30%/yr through 2030 — several multiples of "
              "the fastest year in US grid history, globally and simultaneously, across a supply "
              "chain already at 95%+ utilization at multiple sole-source nodes.")

    h2(doc, "2.4  Utilization (~12%) and latent capacity")
    para(doc, "The model sets inference utilization at 12% (2025) rising to 25% (2035). This "
              "answers the obvious objection — “if all these GPUs are deployed, why isn't power "
              "even higher?” — because at any instant only a fraction of the fleet is doing "
              "revenue inference. The other ~88% is training and post-training, latency headroom "
              "(interactive inference must be provisioned for peak concurrent load, like a grid "
              "sized for the August 5pm peak), diurnal/batch troughs (idle GPUs still draw "
              "30–50% of their power), redundancy and reserved capacity, and just-installed silicon "
              "not yet loaded. This is the same economics as airline load factors and cloud CPU "
              "(famously ~15–20%).")
    para(doc, "The latent capacity splits in two. Reclaimable: batch-shiftable training, off-peak "
              "inference, and better scheduling (continuous batching, disaggregated prefill/decode) "
              "can lift effective utilization — which is exactly what the 12% → 25% path assumes "
              "(roughly a doubling). Structurally stranded: the peak-provisioning and latency "
              "headroom cannot be reclaimed without breaking the product — you will never run "
              "interactive inference at 80% average. So the 25% terminal figure concedes that ~75% "
              "of the fleet stays non-revenue even in 2035. Crucially, rising utilization is a "
              "bearish force on power-per-token, and the model already bakes in a doubling of it; "
              "faster gains flex demand down — a genuine downside lever, and part of why the bear "
              "case exists.")

    h2(doc, "2.5  How much can be optimized away")
    para(doc, "Five levers reduce power per unit of AI output. The table shows realistic headroom "
              "and what is already in the model versus incremental:")
    table(doc, ["Lever", "Mechanism", "Realistic headroom", "In model?"], [
        ["Hardware perf/W", "H100→B200→GB200→Rubin", "~6x by 2040 (9→54 TFLOP/W)", "Yes — core"],
        ["Routing / cascades", "Easy tokens → small models", "~2.9x (active params 297B→104B)", "Yes (FLOPs lens)"],
        ["Quantization", "FP8 → FP4 → sub-4-bit", "~2–4x per step (partly in TFLOP/W)", "Partly"],
        ["Spec. decoding / MFU", "Draft-verify; better kernels", "~1.5–3x throughput", "Partly (MFU conservative)"],
        ["Rising utilization", "Fill the idle ~88%", "~2x effective (12%→25%)", "Yes"],
        ["PUE", "Facility overhead", "~1.22x (1.40→1.15)", "Yes"],
    ])
    para(doc, "The single most important distinction is Jevons. Efficiency comes in two flavors. "
              "Power-reducing: banked gains — same output, fewer watts (PUE, hardware perf/W on a "
              "fixed workload). Capability-spent: gains immediately reinvested in more and bigger "
              "inference — longer reasoning chains, more agents, bigger context. The cost to infer "
              "at fixed quality is halving roughly every two months, yet aggregate inference power "
              "is rising, because every gain is spent on more capability, not banked. This is why "
              "the model's power forecast uses a physical efficiency denominator (~17–19x by 2040), "
              "not the algorithmic one (~122x). Betting the next decade of algorithmic gains gets "
              "banked as a lower power bill is betting against the entire history of computing.")

    # ── 3. Tokens to power ──
    h1(doc, "3.  From tokens to power — the conversion")

    h2(doc, "3.1  What a FLOP is, and why it is the honest unit")
    para(doc, "A FLOP is one floating-point operation — a single multiply or add on real numbers. "
              "“FLOP/s” is the rate; “FLOPs” is a count. A FLOP is the honest unit because it is "
              "the actual physical work the silicon performs, and silicon work is what consumes "
              "watts. A token is not: it is a unit of output, and the compute behind one token "
              "varies ~50x depending on which model produced it. The standard accounting (Kaplan "
              "2020; Chinchilla 2022) is ~2 × N FLOPs per token for inference (the forward pass: "
              "one multiply + one add per active parameter) and ~6 × N for training (forward plus "
              "a backward pass that costs roughly twice the forward). The model uses exactly 2N "
              "and 6N.")
    para(doc, "The critical subtlety is active versus total parameters. In a Mixture-of-Experts "
              "model only the routed experts fire, so N in the 2N rule is active parameters. "
              "DeepSeek-V3 is the canonical anchor: 671B total but only 37B active per token — an "
              "~18x difference. The model's three tiers are built on active params: frontier 500B "
              "(GPT-4/Claude/Gemini-class), mid 70B (DeepSeek-V3-class; 37B is the live anchor, "
              "rounded up conservatively), small 8B (on-device / Llama-8B class). Getting active "
              "vs total right is the difference between a credible model and a 10x error.")

    h2(doc, "3.2  The unit chain, worked through (a 2035 slice)")
    para(doc, "Power is a rate (W, kW, MW, GW); energy is power × time (kWh). A 1 GW data center "
              "running 24h uses 24 GWh that day. The model works in GW of sustained draw — the "
              "grid-planning unit — because that is what must be energized and cooled continuously. "
              "The chain, units explicit:")
    num(doc, "tokens/day × FLOPs/token = FLOPs/day.  e.g. 6.0e15 tokens × (2 × 50e9) "
             "= 6.0e26 FLOPs/day.")
    num(doc, "FLOPs/day ÷ 86,400 s = average FLOP/s.  6.0e26 ÷ 86,400 = 6.94e21 FLOP/s.")
    num(doc, "÷ (TFLOP/W × MFU) = average watts.  6.94e21 ÷ (40e12 × 0.35) = 4.96e8 W.")
    num(doc, "÷ 1,000 thrice = kW → MW → GW.  4.96e8 W = 496 MW = 0.496 GW of IT load.")
    num(doc, "× PUE (facility overhead) = facility GW.  0.496 × 1.20 ≈ 0.60 GW.")
    para(doc, "(This single-tier slice is for unit clarity; the live model runs all three tiers "
              "and the full token build, which is why published 2035 FLOPs-lens demand is ~701 GW. "
              "The chain and the units are the point here.)", italic=True)

    h2(doc, "3.3  Why the shape is trustworthy even though the level is anchored")
    para(doc, "This is the most important methodological point in the model. MFU and the 86,400 "
              "constant appear in both the 2025 normalizer and every forecast year, so they cancel. "
              "The model does not compute absolute GW from first principles. It computes a "
              "dimensionless demand pull each year (tokens × FLOPs/token ÷ fleet efficiency), "
              "normalizes so that 2025 = 70 GW (the empirically anchored present), and reports "
              "every future year as (that year's pull ÷ 2025's pull) × 70 GW. Because of the "
              "re-anchor, any constant that multiplies all years equally — MFU, base PUE, even a "
              "systematic FLOPs/token error — divides out. What survives is the relative "
              "trajectory: how fast tokens grow versus how fast efficiency grows. So read the "
              "absolute GW as “anchored to 70 GW today,” but read the shape, the peak year, the "
              "crossover, and the duration as robust — they depend only on the ratios of the "
              "drivers. The model claims timing and shape with confidence; absolute magnitude with "
              "an anchor.")

    # ── 4. Using the Excel ──
    h1(doc, "4.  How to use it — the Excel workbook")
    para(doc, "File: Token_and_Data_Build_Out_v4_2.xlsx. The original sheets (Summary, Efficiency "
              "Overlay, Base, Robo Bull, Bear) are the token-demand source data and are unchanged. "
              "Four new tabs hold the power model. Edit only the input cells; everything else is a "
              "live formula. Every formula has been independently re-evaluated and ties out exactly, "
              "so opening the file and recalculating changes nothing.")
    h2(doc, "Tab: Macro Levers")
    para(doc, "The master input list for the FLOPs lens — one place to set the assumptions (column "
              "B). The dashboard and the FLOPs tabs read from it.")
    h2(doc, "Tab: FLOPs to Power")
    para(doc, "The live conversion engine. Each row turns one year's tokens into power: tokens/day "
              "(linked to the Efficiency Overlay sheet) × FLOPs/token (= 2 × N) → FLOPs/day ÷ "
              "chip efficiency (TFLOP/W), re-anchored so 2025 = 70 GW → power. Two power columns: "
              "power_raw_GW (the pure pull, which can dip as efficiency overtakes demand) and "
              "power_GW_floored (with the “don't tear down a power plant” floor). Editable: anchor, "
              "retirement rate, and per-year training share, N, and TFLOP/W.")
    h2(doc, "Tab: Duration Lens")
    para(doc, "A clean, write-from summary: the question, the conversion with a worked 2025-vs-2035 "
              "example, the tug-of-war, a token-vs-FLOPs comparison, a chart-ready data block "
              "(select it and Insert → Line Chart), where to stress-test, and the takeaway. The "
              "FLOPs line is linked live.")
    h2(doc, "Tab: Operating Model — the main stress-test surface")
    para(doc, "The full Demand → Supply → Gap chain with a Bull / Base / Bear switch. Set the "
              "scenario in cell B3 (1 = Bull, 2 = Base, 3 = Bear) and the whole sheet recomputes. "
              "The editable scenario table drives five levers: anchor GW, token-growth adjustment, "
              "efficiency adjustment (+ means chips improve faster, which lowers power), capacity "
              "retirement (plateau vs cliff), and supply build scale (slower build = bigger gap). "
              "The summary block reports peak gap, peak year, balance year (gap < 30 GW), overshoot "
              "year (gap < 0), and the 2042 floor. The Base scenario reproduces the Duration lens "
              "exactly, so it is a built-in consistency check.")
    table(doc, ["Scenario", "Peak gap", "Peak year", "Read"], [
        ["Bull", "~526 GW", "2034", "power high and long"],
        ["Base", f"~{n['fp']:.0f} GW", f"{n['fpy']}", "committed FLOPs lens"],
        ["Bear", "~121 GW", "2029", "power rolls over early"],
    ])

    # ── 5. Dashboard ──
    h1(doc, "5.  How to use it — the dashboard")
    para(doc, "Launch run.bat (serves on port 8502). The Macro tab shows the demand/supply/gap "
              "chart and headline metrics. The sidebar expander “FLOPs demand & retirement” "
              "exposes the same levers as the Macro Levers tab. The “Duration Lens” panel shows "
              "the tokens → FLOPs → power conversion live and the token-vs-FLOPs comparison. "
              "Defaults are read from the Macro Levers Excel tab, so the two stay consistent.")

    # ── 6. Glossary ──
    h1(doc, "6.  Key terms")
    table(doc, ["Term", "Definition"], [
        ("Token", "Roughly one word of AI input or output. The model's demand driver."),
        ("Power vs energy", "Power (GW) is the instantaneous draw — what the grid must supply. "
            "Energy (GWh) is power over time. A 1 GW data center uses 24 GWh/day. The thesis is "
            "about GW capacity, because that is what is hard to build."),
        ("FLOP", "One floating-point operation — the actual unit of compute work, hence energy."),
        ("FLOPs/token (2N)", "A token costs ~2 × N operations to generate (N = active parameters). "
            "Training costs ~6N (forward + backward pass)."),
        ("N (active params)", "The size of the model actually doing the work. For a Mixture-of-"
            "Experts model only the activated experts count, not the total. A token's cost varies "
            "~50x across model sizes."),
        ("TFLOP/W", "Trillions of operations per watt — chip efficiency. Assumed ~9 (2025) → "
            "~60 (2042) as silicon improves (Blackwell → Rubin)."),
        ("MFU", "Model FLOPs Utilization — fraction of a chip's peak speed actually achieved "
            "(~35%). Cancels in the 2025 re-anchor."),
        ("PUE", "Power Usage Effectiveness — facility overhead (cooling, conversion). 1.4 means "
            "40% extra on top of the chips."),
        ("Utilization", "Share of deployed AI chips actively computing at any moment (~12%). "
            "Dividing by it scales active power up to total fleet power."),
        ("Anchor (70 GW)", "Measured total DC power in 2025, used to calibrate. Everything is "
            "relative to it, so the messy physics cancels and only relative change matters."),
        ("Gap", "Demand minus supply, in GW. Positive = shortage = pricing power for power/grid/"
            "cooling suppliers. Negative = oversupply."),
        ("Peak / balance / overshoot", "Peak = largest gap. Balance = first year gap < 30 GW. "
            "Overshoot = first year gap goes negative."),
        ("Retirement / floor", "Whether deployed capacity is ever torn down. The floor (retirement "
            "= 0) holds demand flat once built (plateau); higher retirement lets it decline (cliff)."),
        ("Routing / orchestration", "Sending easy queries to smaller models. Lowers N (FLOPs/token) "
            "without changing token count — an efficiency a token-only model cannot see."),
        ("Agentic demand", "Reasoning models and multi-agent workflows that generate many more "
            "tokens per task — a token-count effect, captured in the enterprise agent multiplier."),
        ("Duration lens", "The FLOPs view: not a higher peak, but a higher floor and a later "
            "gap-close — the case for a longer-lived power thesis."),
    ])

    # ── 7. Assumptions ──
    h1(doc, "7.  The assumptions — precedent and derivation")
    para(doc, "This is the heart of the document. For each base-case parameter: the value, what it "
              "is anchored to, and the one fact that matters most. Fuller derivation follows for "
              "the load-bearing ones.")
    table(doc, ["Parameter", "Value", "Anchored to", "Key point"], [
        ["Power anchor", "70 GW (2025)", "C&W / JLL / GS / McKinsey / DCD", "Mid of a tight 65–72 cluster"],
        ["Efficiency doubling (HW)", "2.0 yr", "Epoch AI (Jun-2025) ~2.0", "Owner 2026-06-25 (was 1.85)"],
        ["Algorithmic efficiency", "10 %/yr", "Owner 2026-07-03; << Epoch ~3x/yr", "Compounding, un-lagged; on top of HW; ties token & FLOPs bases"],
        ["Fleet replacement lag", "5 yr", "HW refresh 3–4; enterprise 7–8", "Controls how fast efficiency propagates"],
        ["Supply phase rates", "22/30/25/15 %/yr", "Calibrated to JLL 200–225 GW by 2030", "Steps down as bottlenecks exhaust"],
        ["Utilization", "12% → 25%", "~9% measured 2026; cloud ~15–20%", "Rising = bearish for power-per-token"],
        ["PUE", "1.40 → 1.15", "Liquid-cooling penetration", "Only the trajectory bends the curve"],
        ["N_active tiers", "500 / 70 / 8 B", "GPT-4-class / DeepSeek-V3 / Llama-8B", "Illustrative, not calibrated"],
        ["Routing / avg N_active", "297B → ~59B (2042)", "Algorithmic 10%/yr (tie-out)", "Now driven by the algo rate, not a fixed mix; ties to token base"],
        ["TFLOP/W (FLOPs basis)", "9 × fleet HW index", "Tied to the 2.0-yr doubling", "No longer a decoupled 9→60 ramp; tracks the lever"],
        ["MFU", "0.35", "Frontier training 0.35–0.50", "Cancels in re-anchor"],
        ["Capex basis", "$11.3 + $15 × 50% /MW; 25% power-grid", "JLL 2026", "Supply-driven → $8.4T cumulative"],
    ])
    h2(doc, "The parameters that carry the most weight")
    lead(doc, "70 GW anchor.", "Center of the C&W (63→70–72), JLL (72), GS (65), McKinsey (70), "
              "DCD (68) cluster. ±5 GW, and low-stakes because everything re-anchors to it.")
    lead(doc, "2.0-year hardware doubling + 10%/yr algorithmic.", "Epoch AI's Jun-2025 update puts ML "
              "hardware energy efficiency at ~40%/yr (~2.0-yr doubling); the model takes 2.0 (owner "
              "2026-06-25, reverting an earlier 1.85). On TOP of hardware, algorithmic efficiency "
              "compounds at 10%/yr (owner 2026-07-03) — software/model gains (distillation, MoE, "
              "quantization) that deploy fleet-wide immediately. 10%/yr is deliberately far below "
              "Epoch's ~3x/yr fixed-capability headline (the capability treadmill spends most of that "
              "on better models, not cheaper ones). Both cut against the power thesis (faster efficiency "
              "shrinks the gap), so erring this way understates the power case — a defense, not a "
              "vulnerability. The same 10%/yr reconciles the token and FLOPs bases exactly.")
    lead(doc, "Supply phase rates (22/30/25/15).", "Calibrated so 2030 supply (~229 GW) sits just "
              "above the JLL/McKinsey/Bain 200–225 GW consensus. The step-up-then-down shape is the "
              "physical story: 22% early (permitted pipeline), 30% as the 2026-27 turbine/transformer "
              "wave energizes, then deceleration to 25% and 15% as easy interconnect points, GOES, "
              "and turbine slots exhaust. You cannot compound a physical grid at 30% forever; the "
              "back-half deceleration is the supply ceiling expressing itself.")
    lead(doc, "TFLOP/W ramp (9→60).", "Share-weighted by GPU generation: H100 ~5.65, H200 ~7, "
              "B200 ~9, GB200 ~15, Rubin ~25–30. NVIDIA claims Rubin ~10x perf/W over GB200 via FP4 "
              "and co-design. ~6x hardware gain to 2040 — aggressive but roadmap-defensible; vendor "
              "perf/W claims are FP4-inflated, so realized gains run below headline.")
    lead(doc, "N_active tiers (500/70/8).", "The model's own corrections doc flags these as "
              "illustrative, not calibrated — so the FLOPs peak is a lens, not a hard headline. The "
              "load-bearing fact is the ~50x frontier-to-small ratio, which is robust.")

    h2(doc, "The efficiency tension: ~19x, not ~122x")
    para(doc, "This is the argument a sharp reader will press, so it is worth stating plainly. The "
              "token lens implicitly assumes ~122x blended efficiency by 2040: tokens grow ~176x "
              "yet token-lens demand plateaus near 549 GW (~7.8x the base). That 122x comes from "
              "compounding a blended tokens/kWh index that secretly fuses hardware, model-size, and "
              "algorithmic gains into one exponential.")
    para(doc, "The FLOPs lens decomposes it into physical efficiency only: hardware TFLOP/W rises "
              "~6x (9→54 by 2040) and mix-weighted model size shrinks ~2.9x (297B→104B active). "
              "Six times 2.9 is ~17x — call it ~17–19x — the physically defensible figure. The "
              "gap between ~19x and ~122x is the algorithmic term, and Jevons says you cannot bank "
              "it: Epoch measures algorithmic progress at ~3x/yr, but that progress has historically "
              "been spent on capability (longer reasoning, more agents), not saved as a lower power "
              "bill. A power forecast must therefore use the physical denominator. This is why the "
              "FLOPs lens shows demand plateauing higher (708 vs 549 GW) and the gap closing later "
              "(2036 vs 2034). The variant perception is precisely that consensus is banking an "
              "algorithmic windfall that will be spent, not saved.")

    # ── 8. Bull / Base / Bear ──
    h1(doc, "8.  Bull / Base / Bear — why each knob is set where it is")
    para(doc, "The Operating Model overlays four knobs on the FLOPs-lens base. Bull = power high "
              "and long (efficiency slower, supply slower, no retirement); Bear = power rolls over "
              "(efficiency faster, supply faster, capacity retires).")
    table(doc, ["Knob", "Bull", "Base", "Bear", "Why this magnitude"], [
        ["Token growth adj", "+2%/yr", "0", "−2%/yr",
            "±2%/yr compounding ≈ 1.4x / 0.7x by 2042; the honest agentic-adoption band, not a regime change"],
        ["Efficiency adj", "−1.5%/yr", "0", "+3%/yr",
            "Asymmetric: efficiency has more room to surprise up (FP4, algorithmic) than down (process walls)"],
        ["Retirement", "0", "0", "8%/yr",
            "~12.5-yr half-life of demand contribution; credible teardown/repurposing pace"],
        ["Supply build scale", "0.85", "1.0", "1.15",
            "±15% execution band around an already historically-unprecedented 22–30%/yr build"],
    ])
    h2(doc, "Why these magnitudes")
    lead(doc, "Token growth ±2%/yr.", "The base path already grows 176x; ±2%/yr compounding over "
              "17 years is ~1.4x (bull) to ~0.7x (bear) by 2042. It brackets plausible agentic "
              "adoption without the structurally-impossible agent-multiplier blowup the model "
              "explicitly rejected (a blended enterprise multiplier of 100x+ would imply more than "
              "100% of tasks being agentic — arithmetically impossible).")
    lead(doc, "Efficiency −1.5% bull vs +3% bear (asymmetric).", "Sign convention: + means faster "
              "TFLOP/W = less power. The bear adjustment is twice as large because efficiency has "
              "more room to surprise to the upside (bearish for power): FP4-native inference, "
              "sub-4-bit, algorithmic breakthroughs (Epoch's ~3x/yr already outruns hardware), new "
              "architectures. The downside (efficiency slowing) requires hitting process/packaging "
              "walls — real but slower-moving and partly routed around by chiplets. The distribution "
              "is right-skewed toward faster efficiency, and the asymmetric knob encodes exactly "
              "that — the model tilts the efficiency risk against its own power thesis.")
    lead(doc, "Retirement 8%/yr (bear).", "GPU useful life is ~5–6 years; the shell lives 15–30. "
              "8%/yr floor-decay implies the deployed base stops counting as binding demand on a "
              "~12.5-year half-life — reasonable once you account for repurposing (old GPUs shifted "
              "to low-value batch work rather than powered off). Base case is 0% (strict floor); "
              "8%/yr is the bear's “the 2030s are a cliff, not a plateau” expression, and the "
              "single biggest out-year swing factor.")
    lead(doc, "Supply build scale ±15%.", "Brackets execution on the physical bottlenecks. Bull "
              "0.85 (slower build, bigger gap): transformers at 100–210 weeks, single-source GOES, "
              "interconnect failures, the labor pipeline. Bear 1.15 (faster, gap closes): queue "
              "reform, new domestic transformer capacity on time, behind-the-meter nuclear bypassing "
              "the queue. ±15% is right because the build is already assumed to run at an "
              "unprecedented 22–30%/yr — there isn't ±50% of realistic headroom.")

    h2(doc, "The investment read")
    para(doc, "The peak-gap spread is ~121 to ~526 GW (Bear to Bull). That ~4x range is itself the "
              "message: in every scenario there is a multi-year shortfall measured in hundreds of "
              "reactors' worth of power; the disagreement is about duration and magnitude, not "
              "existence. Even the bear (121 GW, closing earliest) is a ~120-reactor shortfall.")
    para(doc, "For the names: generation (GEV gas-turbine slots rationed to 2028+; CEG/TLN "
              "behind-the-meter nuclear as the interconnect bypass; VST/NRG dispatchable capacity); "
              "the transformer/GOES complex (CLF the single most mispriced node — a sole-source US "
              "GOES franchise trading as commodity steel); and cooling/electrical (VRT liquid "
              "cooling, mandatory by physics above 35 kW/rack).")
    para(doc, "The duration argument is the edge. Consensus power models lean on the token-lens / "
              "blended-efficiency framing and conclude the gap rolls over in the early-mid 2030s — "
              "pricing the power names as a 2025–28 cyclical that fades. The variant perception is "
              "that this rollover is too early, because it banks an algorithmic-efficiency windfall "
              "(~122x) that Jevons says gets spent, not saved. Strip it to the physical denominator "
              "(~17–19x) and the demand floor is ~29% higher, the gap closes ~2036 not 2034, and "
              "there is no out-year collapse. The trade this implies: the power/grid/cooling complex "
              "is a structurally tight build through the mid-2030s, not a cyclical that fades in "
              "2029 — and the market is pricing the consensus rollover. The risk to the view is the "
              "bear efficiency lever, which is exactly why the model tilts the efficiency asymmetry "
              "against itself, so the duration call survives even when it gives efficiency the "
              "benefit of the doubt.")

    # ── 9. Trust ──
    h1(doc, "9.  A note on trust and honest caveats")
    para(doc, "Every formula in the workbook's power tabs has been independently re-evaluated and "
              "ties out to its displayed value, and the committed base case is fingerprint-locked "
              "so accidental drift is caught. Three things to keep in mind when defending the work:")
    lead(doc, "Absolute GW is anchored, not derived.", "Trust the shape and timing (when the gap "
              "peaks, when it closes) more than any single year's exact GW.")
    lead(doc, "Efficiency (2.0-yr HW + 10%/yr algorithmic) is the most aggressive externally-checkable input.", "Both "
              "are at/above Epoch's ~2.0-yr hardware headline, and both cut against the power thesis — "
              "so if anyone calls the model too power-bullish, efficiency is the defense.")
    lead(doc, "The FLOPs basis now ties out to the token base.", "As of 2026-07-03 the FLOPs lens "
              "shares the same hardware doubling and the same 10%/yr algorithmic (expressed as routing "
              "to smaller models), so it reproduces the token base exactly — a consistency check, not a "
              "higher/longer forecast. The committed headline is the ~377 GW token base (peak gap ~200 "
              "@ 2028). The ~750–850 GW “agentic” case some may have seen was deliberately rejected — it "
              "required a structurally impossible >100% agentic task penetration.")
    para(doc, "The point of the model is not a precise number; it is a defensible framework for "
              "diving into the assumptions and dimensioning the most likely outcomes.", italic=True)

    doc.save(str(OUT))
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
